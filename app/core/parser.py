"""文档解析器 — PyMuPDF 逐页分类 + EasyOCR + 表格展平 + 语义感知"""

import re
import io
import logging
import hashlib
from pathlib import Path
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff", ".tif"}


@dataclass
class ParsedPage:
    """解析后的单页"""
    page_num: int
    text: str
    page_type: str = "digital"   # digital | scanned | mixed
    ocr_confidence: float = 0.0  # 仅 scanned/mixed 有效
    has_table: bool = False
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """解析后的完整文档"""
    pages: list[ParsedPage]
    full_text: str
    metadata: dict
    content_hash: str


class DocumentParser:
    """统一文档解析：PDF(分类)/图片(OCR)/文本"""

    # ── 公共入口 ──

    @staticmethod
    def parse(file_path: str) -> tuple[str, dict]:
        """解析文件 → (文本内容, 元数据)"""
        path = Path(file_path)
        suffix = path.suffix.lower()
        content = ""
        meta_extra = {}

        if suffix == ".pdf":
            content, meta_extra = DocumentParser._parse_pdf_classified(path)
        elif suffix in IMAGE_SUFFIXES:
            content, meta_extra = DocumentParser._parse_image_ocr(path)
        elif suffix in (".docx",):
            content, meta_extra = DocumentParser._parse_docx(path)
        elif suffix in (".xlsx", ".xls"):
            content, meta_extra = DocumentParser._parse_excel(path)
        else:
            content = DocumentParser._parse_text(path)

        # 基础元数据
        metadata = {
            "file_name": path.name,
            "file_path": str(path.absolute()),
            "file_type": suffix.lstrip("."),
            "file_size": path.stat().st_size,
            **meta_extra,
        }
        return content, metadata

    # ── PDF: PyMuPDF 逐页分类 ──

    @staticmethod
    def _parse_pdf_classified(path: Path) -> tuple[str, dict]:
        """用 PyMuPDF 逐页分类处理"""
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        parsed_pages = []
        meta_extra = {
            "total_pages": len(doc),
            "digital_pages": 0,
            "scanned_pages": 0,
            "mixed_pages": 0,
            "has_ocr": False,
            "page_types": [],
        }

        for i, page in enumerate(doc):
            page_num = i + 1
            page_text = page.get_text("text").strip()
            text_len = len(page_text)

            # 分类：基于文字量判断页面类型
            if text_len > 100:
                page_type = "digital"
                final_text = page_text
                ocr_conf = 0.0
            elif text_len > 20:
                page_type = "mixed"
                # 混合页：提取文字 + OCR 补充
                ocr_text = DocumentParser._ocr_page_image(page, path.name, page_num)
                final_text = page_text + "\n[OCR补充]\n" + ocr_text if ocr_text.strip() else page_text
                ocr_conf = 0.5
            else:
                page_type = "scanned"
                final_text = DocumentParser._ocr_page_image(page, path.name, page_num)
                ocr_conf = 0.7  # OCR 是主来源
                if not final_text.strip():
                    final_text = f"[第{page_num}页: 空白或不可识别]"

            # 检测表格
            has_table = DocumentParser._page_has_table(final_text)

            # 展平表格
            if has_table:
                final_text = DocumentParser._flatten_tables(final_text)

            parsed_pages.append(ParsedPage(
                page_num=page_num,
                text=final_text,
                page_type=page_type,
                ocr_confidence=ocr_conf,
                has_table=has_table,
                metadata={"page": page_num, "type": page_type},
            ))

            # 统计
            if page_type == "digital":
                meta_extra["digital_pages"] += 1
            elif page_type == "scanned":
                meta_extra["scanned_pages"] += 1
                meta_extra["has_ocr"] = True
            else:
                meta_extra["mixed_pages"] += 1
                meta_extra["has_ocr"] = True

            meta_extra["page_types"].append(page_type)

        total_pages = len(doc)
        doc.close()

        # 组装全文
        full_text_parts = []
        for p in parsed_pages:
            header = f"[第 {p.page_num} 页]"
            if p.page_type != "digital":
                header += f" (OCR识别, 置信度: {p.ocr_confidence:.0%})"
            if p.has_table:
                header += " [含表格]"
            full_text_parts.append(f"{header}\n{p.text}")

        full_text = "\n\n".join(full_text_parts)
        content_hash = hashlib.md5(full_text.encode()).hexdigest()

        logger.info(
            "PDF parsed: %s → %d pages (digital=%d, scanned=%d, mixed=%d), %d chars",
            path.name, total_pages, meta_extra["digital_pages"],
            meta_extra["scanned_pages"], meta_extra["mixed_pages"], len(full_text),
        )

        meta_extra["content_hash"] = content_hash
        return full_text, meta_extra

    # ── 图片 OCR ──

    @staticmethod
    def _parse_image_ocr(path: Path) -> tuple[str, dict]:
        """用 EasyOCR 提取图片文字"""
        from app.core.vision import VisionService
        vision = VisionService()
        text = vision.describe_image(str(path))
        content_hash = hashlib.md5(text.encode()).hexdigest()
        logger.info("Image OCR: %s → %d chars", path.name, len(text))
        return text, {
            "content_hash": content_hash,
            "page_type": "image_ocr",
            "ocr_confidence": 0.8,
        }

    # ── DOCX ──

    @staticmethod
    def _parse_docx(path: Path) -> tuple[str, dict]:
        """解析 Word 文档"""
        try:
            import docx
            d = docx.Document(str(path))
            paragraphs = [p.text for p in d.paragraphs if p.text.strip()]
            # 也提取表格
            for table in d.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                paragraphs.append("\n".join(rows))
            text = "\n\n".join(paragraphs)
        except ImportError:
            text = DocumentParser._parse_text(path)
        content_hash = hashlib.md5(text.encode()).hexdigest()
        logger.info("DOCX parsed: %s → %d chars", path.name, len(text))
        return text, {"content_hash": content_hash}

    # ── Excel ──

    @staticmethod
    def _parse_excel(path: Path) -> tuple[str, dict]:
        """解析 Excel 为 Markdown 表格"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"## Sheet: {sheet_name}")
                rows = list(ws.iter_rows(values_only=True))
                for row in rows[:500]:  # 最多 500 行
                    parts.append(" | ".join(str(c) if c is not None else "" for c in row))
                if len(rows) > 500:
                    parts.append(f"... (共 {len(rows)} 行，仅展示前 500 行)")
            wb.close()
            text = "\n".join(parts)
        except ImportError:
            text = DocumentParser._parse_text(path)
        content_hash = hashlib.md5(text.encode()).hexdigest()
        logger.info("Excel parsed: %s → %d chars", path.name, len(text))
        return text, {"content_hash": content_hash}

    # ── 纯文本 ──

    @staticmethod
    def _parse_text(path: Path) -> str:
        """解析纯文本/代码文件"""
        try:
            with open(path, "r", encoding="utf-8") as f:
                raw = f.read()
        except UnicodeDecodeError:
            with open(path, "r", encoding="latin-1") as f:
                raw = f.read()

        suffix = path.suffix.lower()
        if suffix == ".md":
            return raw
        elif suffix in (".html", ".htm"):
            return DocumentParser._strip_html(raw)
        elif suffix in (".py", ".js", ".ts", ".java", ".go", ".rs", ".sql"):
            return f"```{suffix[1:]}\n{raw}\n```"
        return raw

    # ── OCR 页面图片 ──

    @staticmethod
    def _ocr_page_image(page, filename: str, page_num: int) -> str:
        """对 PDF 单页渲染图片并 OCR"""
        try:
            # 渲染页面为图片
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")

            # EasyOCR
            from app.core.vision import _get_ocr_reader
            reader = _get_ocr_reader()
            import numpy as np
            from PIL import Image
            img = Image.open(io.BytesIO(img_bytes))
            img_np = np.array(img)
            results = reader.readtext(img_np, detail=0)
            if results:
                return "\n".join(results)
            return ""
        except Exception as e:
            logger.debug("OCR failed for %s page %d: %s", filename, page_num, e)
            return ""

    # ── 表格检测与展平 ──

    @staticmethod
    def _page_has_table(text: str) -> bool:
        """简单启发式检测表格：连续的 | 分隔符或对齐的空白"""
        lines = text.split("\n")
        pipe_lines = sum(1 for l in lines if l.count("|") >= 2)
        if pipe_lines >= 3:
            return True
        # 检测制表符分隔
        tab_lines = sum(1 for l in lines if "\t" in l)
        return tab_lines >= 3

    @staticmethod
    def _flatten_tables(text: str) -> str:
        """将表格展平为可读的 header-row 格式，方便 BM25 匹配"""
        lines = text.split("\n")
        result = []
        in_table = False
        table_rows = []

        for line in lines:
            if line.count("|") >= 2 or "\t" in line:
                in_table = True
                sep = "\t" if "\t" in line else "|"
                cells = [c.strip() for c in line.split(sep) if c.strip()]
                table_rows.append(cells)
            else:
                if in_table and table_rows:
                    result.append(DocumentParser._format_table_block(table_rows))
                    table_rows = []
                in_table = False
                result.append(line)

        if table_rows:
            result.append(DocumentParser._format_table_block(table_rows))

        return "\n".join(result)

    @staticmethod
    def _format_table_block(rows: list[list[str]]) -> str:
        """将表格行格式化为 header: value 片段"""
        if len(rows) < 2:
            return " | ".join(rows[0]) if rows else ""

        header = rows[0]
        formatted = ["[表格]"]
        for row in rows[1:]:
            parts = []
            for i, cell in enumerate(row):
                if i < len(header) and cell.strip():
                    parts.append(f"{header[i]}: {cell}")
            formatted.append("; ".join(parts))

        return "\n".join(formatted)

    # ── HTML ──

    @staticmethod
    def _strip_html(html: str) -> str:
        text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text
